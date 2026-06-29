"""Academic paper review service powered by Groq and Gemini."""

from __future__ import annotations

import io
import json
import os
import re
from collections import Counter
from pathlib import Path
from urllib.parse import urlparse

import httpx
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pypdf import PdfReader
from docx import Document

from llm_providers import (
    fit_review_content,
    get_limits,
    get_text_limit,
    resolve_provider,
    run_with_fallback,
    trim_context_text,
)

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
STATIC_DIR = BASE_DIR / "static"

MAX_FILE_BYTES = 50 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}
ALLOWED_DRIVE_HOSTS = {"drive.google.com", "docs.google.com"}

SYSTEM_PROMPT = """You are Professor Meridian, an experienced academic advisor and thesis/capstone consultant. Your role is to review student research papers and provide constructive, rigorous feedback.

SCOPE — ACADEMIC ONLY:
- You ONLY discuss academic writing, research methodology, argument structure, literature review, citations, data analysis, results interpretation, academic tone, formatting, and scholarly conventions.
- If the user asks about non-academic topics (entertainment, personal advice, coding help unrelated to the paper, politics, etc.), politely decline and redirect: "I can only assist with academic paper review and scholarly writing."
- If the uploaded document is not academic in nature, explain that and offer feedback only on any writing-quality aspects that apply to formal writing.

REVIEW APPROACH:
- Be supportive but honest, like a good teacher — never harsh, always actionable.
- Ground every critique in specific evidence from the paper (quote or paraphrase sections).
- Distinguish major issues (methodology, thesis clarity, missing evidence) from minor ones (word choice, formatting).
- Suggest concrete revisions, not vague advice.

REFERENCE MATERIAL (OPTIONAL):
- The student may optionally provide Google Drive reference material. It may be missing, empty, or unavailable.
- When no reference material is present, rely on the uploaded document, your earlier review (if any), and the conversation history.
- Do not tell the student that Google Drive is required or that you are missing external reference files unless they ask about that link specifically.

CONSULTATION JOURNAL (OPTIONAL):
- The student may provide a consultation journal documenting conversations with other professors or advisors.
- Use it to understand prior feedback, avoid repeating the same advice, and build on what other mentors already said.
- When no journal is provided, rely on the uploaded paper and conversation history only.

QUESTIONNAIRE (OPTIONAL — REQUIRES PAPER):
- The student may provide a research questionnaire or survey instrument related to the paper.
- Only assess the questionnaire in relation to the uploaded paper's main argument — never review a questionnaire without the paper.
- When questionnaire critique is enabled, include a ## Questionnaire Assessment section.

OUTPUT FORMAT (use markdown headings):
## Overall Assessment
Brief summary of the paper's quality and main contribution.

## Strengths
Bullet points of what works well.

## Areas for Improvement
Organized by section (Introduction, Literature Review, Methodology, Results, Discussion, Conclusion) when applicable.

## Specific Revision Suggestions
Numbered, actionable changes the author should make.

## Academic Writing & Style
Comments on tone, clarity, passive/active voice, paragraph structure, transitions.

## Citations & References
Note any gaps, formatting issues, or missing sources (if references are present).

## Questions for the Author
Thoughtful questions a thesis advisor would ask to deepen the work.

Keep feedback practical and proportionate to the paper's apparent level (undergraduate capstone, graduate thesis, journal article, etc.)."""

COMPACT_SYSTEM_PROMPT = """You are Professor Meridian, an academic thesis advisor. Review student papers with constructive, evidence-based feedback. Academic topics only.

Use markdown headings: Overall Assessment, Strengths, Areas for Improvement, Specific Revision Suggestions, Academic Writing & Style, Citations & References, Questions for the Author.

Be supportive, honest, and actionable. Quote or paraphrase the paper when critiquing. If a consultation journal from other professors is provided, align with their feedback and avoid repeating the same advice."""

QUESTIONNAIRE_CRITIQUE_PROMPT = """QUESTIONNAIRE CRITIQUE (ENABLED):
Include a dedicated section ## Questionnaire Assessment in your review.
- State the paper's main argument or thesis claim clearly.
- Evaluate whether each major part of the questionnaire supports, operationalizes, or helps answer that argument.
- Flag missing, vague, leading, or misaligned questions; suggest concrete improvements.
- Note gaps where the questionnaire fails to gather evidence the paper needs, or duplicates claims without adding measurable support."""

QUESTIONNAIRE_ALIGNMENT_PROMPT = """QUESTIONNAIRE ALIGNMENT (INCLUDED):
Briefly note in your review whether the questionnaire aligns with and supports the paper's main argument. If misalignment is significant, say so even without a full questionnaire critique section."""

QUESTIONING_INSTRUCTIONS: dict[str, dict[str, str]] = {
    "light": {
        "review": """QUESTIONING MODE — LIGHT:
- In "Questions for the Author", include 2–4 questions only where critical information is missing or a major claim lacks support.
- Do not challenge every weak point — prioritize the most important gaps.""",
        "chat": """QUESTIONING MODE — LIGHT:
- Ask a clarifying question only when the student's message is vague or you cannot give useful advice without more detail.
- Do not routinely challenge their claims.""",
    },
    "balanced": {
        "review": """QUESTIONING MODE — BALANCED:
- In "Questions for the Author", ask 4–7 thoughtful advisor questions.
- When you notice unsupported claims, logical gaps, unclear methodology, or inconsistencies, pose direct questions that push the student to explain or justify those parts.
- Briefly flag questionable passages, then ask what evidence or reasoning supports them.""",
        "chat": """QUESTIONING MODE — BALANCED:
- When the student makes a claim about their paper, argument, or research that seems weak, unsupported, or inconsistent with their document, politely question it.
- Ask 1–2 probing questions per reply when you spot genuine gaps in reasoning or evidence.
- Example tone: 'You mention X is established — what source or finding in your paper supports that?'""",
    },
    "proactive": {
        "review": """QUESTIONING MODE — PROACTIVE (Socratic):
- Treat "Questions for the Author" as a rigorous thesis-defense segment — aim for 6–10 sharp, specific questions.
- Actively challenge every questionable claim, unsupported generalization, methodology weakness, citation gap, and logical leap in the paper.
- For each concern, name the passage or issue and ask the student to defend, clarify, or revise it.
- Prepare the student as if they will face a thesis committee.""",
        "chat": """QUESTIONING MODE — PROACTIVE (Socratic):
- Act like a thesis advisor in a mock defense: when anything the student says, or wrote in their paper, is questionable, incomplete, or self-contradictory, question it directly but supportively.
- End most replies with at least one probing question that pushes deeper thinking.
- Challenge assumptions, unsupported claims, and vague reasoning. Ask them to cite evidence, explain methodology choices, or reconcile inconsistencies.
- If their statement conflicts with the paper excerpt or your earlier review, point that out and ask them to explain.""",
    },
}


THESIS_TEMPLATE = """PROPOSED TECH THESIS/CAPSTONE TEMPLATE / FORMAT 2025–2026

THESIS 1 / CAPSTONE 1 — Structure:
Chapter 1: Introduction
  A. Project Rationale
     a. Research Context
        i. Review of Related Literature
        ii. Research Methods and Findings
     b. Statement of the Problem
  B. Project Narrative
     a. Description and Objectives
     b. Scope and Limitations
     c. SWOT Analysis
     d. Project Testing and Evaluation
        i. Respondents
        ii. Instrument
        iii. Analysis (Thematic / Statistical)
Chapter 2: Review of Related Works
Chapter 3: Design Methodology
  A. Conceptual Framework
  B. Project Timeline
  C. System Scope / (For EMC) High Concept Game Design Document
  D. Requirements Specification / (For EMC) Prototype Overview
Bibliography
Appendices
Focus: determine the topic/focus of the output; establish the problem/need and its context through literature review and preliminary data gathering (interview, survey, focus group discussion, observation); conceptualize the project that addresses the problem; determine objectives; identify scope and limitations; conduct SWOT analysis; establish how the project will be evaluated against objectives; review existing related technologies/systems/games/applications; plan the design methodology.

THESIS 2 / CAPSTONE 2 — Structure:
Chapter 3: Design Methodology (cont.) — A. Conceptual Framework, B. Project Timeline, C. System Scope, D. Requirements Specification
Chapter 4: Results and Interpretations
Chapter 5: Conclusions and Next Steps
Bibliography (updated)
Appendices (updated)
Focus: implement the full design methodology; conduct prototype testing; collect and report findings and feedback; draw conclusions and determine next steps.

SECTION DESCRIPTIONS:
- Chapter 1 Introduction: preliminaries to project development and implementation.
  - A. Project Rationale: justification for the project; presents preliminary research on the topic and zeroes in on the actual problem/issue to be addressed.
    - a. Research Context: establishes that a problem exists that technology can address (made of Review of Related Literature + Research Methods and Findings).
      - i. Review of Related Literature: cite varied sources (news/web articles, studies, government documents) for historical/social/cultural context; ideally written thematically.
      - ii. Research Methods and Findings: preliminary data gathering (expert interviews, immersion, etc.) to deepen understanding and identify technology's role; report key findings.
    - b. Statement of the Problem: the problem/issue/gap to be addressed; stems from the research context.
  - B. Project Narrative: explains the proposed project in light of the identified problem.
    - a. Description and Objectives: describes the project and its specific objectives; objectives must align with the Statement of the Problem and are the basis for testing/evaluation.
    - b. Scope and Limitations: the extent to which the project addresses the problem and meets the objectives.
    - c. SWOT Analysis: strengths, weaknesses, opportunities, threats; also identifies marketability / unique selling proposition.
    - d. Project Testing and Evaluation: how the project is tested/evaluated, based on the objectives.
      - i. Respondents: from primary stakeholders / target audience; generally at least 30 respondents for prototype testing and feedback.
      - ii. Instrument: type and content aligned with the objectives (survey questionnaire, test, checklist, etc.).
      - iii. Analysis: thematic analysis for qualitative data (interviews, FGDs); statistical analysis for quantitative data (closed-ended surveys, tests).
- Chapter 2 Review of Related Works: related works/technologies that address the same/similar topic OR have features to include; proponents should gain hands-on experience with them when possible.
- Chapter 3 Design Methodology: the actual development of the project.
- Chapter 4 Results and Interpretations: results from the testing/evaluation phase with the proponents' interpretations. (For EMC) includes the finalized Game Design Document.
- Chapter 5 Conclusions and Next Steps: overall assessment with key insights on the topic and on technology's uses/potentials/limitations; ends with next steps based on user and panel feedback.
- Bibliography: alphabetical listing of sources/works reviewed.
- Appendices: supporting documents for the project.

Reference: Lavina, C.G., Manabo, R.D., Hernandez, G.D.C., Hablanida, F.L., Lacorte, A.M., & Gaza-Ebron, J. (2022). Outcomes-based practical guide to thesis and capstone project writing in computing. Mindshapers Co., Inc."""

THESIS_TEMPLATE_COMPACT = """TECH THESIS/CAPSTONE TEMPLATE 2025–2026 (expected structure):
Thesis/Capstone 1:
- Ch.1 Introduction: A. Project Rationale [a. Research Context (i. Review of Related Literature, ii. Research Methods and Findings); b. Statement of the Problem]; B. Project Narrative [a. Description & Objectives; b. Scope & Limitations; c. SWOT Analysis; d. Project Testing & Evaluation (i. Respondents ~30 min; ii. Instrument; iii. Analysis thematic/statistical)].
- Ch.2 Review of Related Works.
- Ch.3 Design Methodology: A. Conceptual Framework; B. Project Timeline; C. System Scope / (EMC) High Concept GDD; D. Requirements Specification / (EMC) Prototype Overview.
- Bibliography; Appendices.
Thesis/Capstone 2:
- Ch.3 Design Methodology (cont.); Ch.4 Results and Interpretations; Ch.5 Conclusions and Next Steps; updated Bibliography & Appendices.
Key rule: objectives must align with the Statement of the Problem and be the basis for testing/evaluation."""


def normalize_questioning_mode(mode: str) -> str:
    normalized = (mode or "balanced").strip().lower()
    if normalized not in QUESTIONING_INSTRUCTIONS:
        return "balanced"
    return normalized


def get_questioning_prompt(mode: str, context: str) -> str:
    normalized = normalize_questioning_mode(mode)
    return QUESTIONING_INSTRUCTIONS[normalized][context]


def get_template_block(provider: str) -> str:
    limits = get_limits(provider)
    template = THESIS_TEMPLATE_COMPACT if limits.use_compact_prompt else THESIS_TEMPLATE
    return (
        "\n\nINSTITUTIONAL THESIS/CAPSTONE TEMPLATE (for reference):\n"
        "Use this official 2025–2026 template as the expected structure and format when reviewing a "
        "Thesis, Capstone, or Feasibility Study. Check whether the paper follows these chapters and "
        "sections, and clearly flag any missing, incomplete, or out-of-order parts. For other paper "
        "types, apply it only where relevant.\n"
        "--- BEGIN TEMPLATE ---\n"
        f"{template}\n"
        "--- END TEMPLATE ---"
    )


def should_include_template(paper_type: str) -> bool:
    """The institutional thesis/capstone template does not apply to feasibility studies."""
    return "feasibility" not in (paper_type or "").strip().lower()


def get_base_system_prompt(provider: str, include_template: bool = True) -> str:
    limits = get_limits(provider)
    base = COMPACT_SYSTEM_PROMPT if limits.use_compact_prompt else SYSTEM_PROMPT
    if include_template:
        base += get_template_block(provider)
    return base


def get_review_system_prompt(
    provider: str,
    questioning_mode: str,
    critique_questionnaire: bool = False,
    include_questionnaire: bool = False,
    paper_type: str = "",
) -> str:
    prompt = (
        get_base_system_prompt(provider, include_template=should_include_template(paper_type))
        + "\n\n"
        + get_questioning_prompt(questioning_mode, "review")
    )
    if critique_questionnaire:
        prompt += "\n\n" + QUESTIONNAIRE_CRITIQUE_PROMPT
    elif include_questionnaire:
        prompt += "\n\n" + QUESTIONNAIRE_ALIGNMENT_PROMPT
    return prompt


def parse_form_bool(value: str) -> bool:
    return (value or "").strip().lower() in {"1", "true", "yes", "on"}


def validate_drive_url(url: str) -> str:
    cleaned = url.strip()
    if not cleaned:
        return ""
    parsed = urlparse(cleaned)
    if parsed.scheme not in {"http", "https"}:
        raise HTTPException(status_code=400, detail="Invalid link. Use an http or https Google Drive URL.")
    host = parsed.netloc.lower().removeprefix("www.")
    if host not in ALLOWED_DRIVE_HOSTS:
        raise HTTPException(
            status_code=400,
            detail="Only Google Drive and Google Docs links are supported.",
        )
    return cleaned


def extract_drive_file_id(url: str) -> str | None:
    patterns = [
        r"/file/d/([a-zA-Z0-9_-]+)",
        r"/document/d/([a-zA-Z0-9_-]+)",
        r"/spreadsheets/d/([a-zA-Z0-9_-]+)",
        r"/presentation/d/([a-zA-Z0-9_-]+)",
        r"[?&]id=([a-zA-Z0-9_-]+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None


def classify_drive_link(url: str) -> str:
    if "/document/" in url:
        return "document"
    if "/spreadsheets/" in url:
        return "spreadsheet"
    if "/presentation/" in url:
        return "presentation"
    return "file"


def looks_like_html(data: bytes) -> bool:
    sample = data[:800].lower()
    return b"<!doctype html" in sample or b"<html" in sample


def bytes_to_reference_text(data: bytes, kind: str) -> str:
    if looks_like_html(data):
        raise HTTPException(
            status_code=400,
            detail=(
                "Could not access that Google Drive file. "
                "Set sharing to 'Anyone with the link can view' and try again."
            ),
        )

    if kind in {"document", "spreadsheet", "presentation"}:
        return data.decode("utf-8", errors="replace")

    if data[:4] == b"%PDF":
        return extract_text_from_pdf(data)

    if data[:2] == b"PK":
        return extract_text_from_docx(data)

    return data.decode("utf-8", errors="replace")


def download_drive_bytes(client: httpx.Client, file_id: str, export_url: str | None = None) -> bytes:
    if export_url:
        response = client.get(export_url)
        response.raise_for_status()
        return response.content

    url = "https://drive.google.com/uc"
    response = client.get(url, params={"export": "download", "id": file_id})
    token = next((value for key, value in response.cookies.items() if key.startswith("download_warning")), None)
    if token:
        response = client.get(url, params={"export": "download", "id": file_id, "confirm": token})
    response.raise_for_status()
    return response.content


def fetch_reference_text(link: str) -> tuple[str, str | None]:
    cleaned = link.strip()
    if not cleaned:
        return "", None

    try:
        cleaned = validate_drive_url(cleaned)
    except HTTPException as exc:
        return "", str(exc.detail)

    file_id = extract_drive_file_id(cleaned)
    if not file_id:
        return "", "Could not read a file ID from that Google Drive link."

    kind = classify_drive_link(cleaned)
    export_urls = {
        "document": f"https://docs.google.com/document/d/{file_id}/export?format=txt",
        "spreadsheet": f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=csv",
        "presentation": f"https://docs.google.com/presentation/d/{file_id}/export?format=txt",
    }

    try:
        with httpx.Client(timeout=30.0, follow_redirects=True) as client:
            if kind in export_urls:
                data = download_drive_bytes(client, file_id, export_urls[kind])
            else:
                data = download_drive_bytes(client, file_id)
    except Exception:
        return "", "Could not access that Google Drive file. Check sharing settings or leave the link empty."

    try:
        text = bytes_to_reference_text(data, kind)
    except HTTPException as exc:
        return "", str(exc.detail)

    text = compress_text(text)
    if not text:
        return "", "That Google Drive file appears to be empty or has no readable text."
    return text, None


def resolve_reference_material(
    reference_link: str,
    reference_excerpt: str = "",
    max_reference_chars: int = 1500,
) -> tuple[str, bool, str | None]:
    excerpt = reference_excerpt.strip()
    if excerpt:
        truncated = len(excerpt) > max_reference_chars
        return excerpt[:max_reference_chars], truncated, None

    link = reference_link.strip()
    if not link:
        return "", False, None

    text, warning = fetch_reference_text(link)
    if not text:
        return "", False, warning

    truncated = len(text) > max_reference_chars
    return text[:max_reference_chars], truncated, None


def build_reference_response(link: str, text: str, truncated: bool, warning: str | None) -> dict:
    word_count = len(text.split()) if text else 0
    return {
        "ok": True,
        "empty": not bool(text),
        "message": warning
        or (
            "No Google Drive reference loaded. The professor will use your uploaded paper and conversation context."
            if not text
            else None
        ),
        "word_count": word_count,
        "truncated": truncated,
        "excerpt": text[:8000],
        "reference_text": text,
        "link": link.strip() if text else "",
    }


def format_reference_block(reference_text: str, source_link: str = "") -> str:
    if not reference_text.strip():
        return ""
    source_note = f"\nSource link: {source_link.strip()}" if source_link.strip() else ""
    return (
        "\n\nAdditional reference material from Google Drive"
        " (use this for context when reviewing or advising on the main document):"
        f"{source_note}\n"
        "--- BEGIN REFERENCE MATERIAL ---\n"
        f"{reference_text.strip()}\n"
        "--- END REFERENCE MATERIAL ---"
    )


def format_journal_block(journal_text: str, filename: str = "") -> str:
    if not journal_text.strip():
        return ""
    source_note = f"\nJournal file: {filename.strip()}" if filename.strip() else ""
    return (
        "\n\nConsultation journal — notes from the student's meetings with other professors/advisors"
        " (use this to align with prior feedback and avoid repeating advice):"
        f"{source_note}\n"
        "--- BEGIN CONSULTATION JOURNAL ---\n"
        f"{journal_text.strip()}\n"
        "--- END CONSULTATION JOURNAL ---"
    )


def format_questionnaire_block(
    questionnaire_text: str,
    filename: str = "",
    critique_enabled: bool = False,
) -> str:
    if not questionnaire_text.strip():
        return ""
    source_note = f"\nQuestionnaire file: {filename.strip()}" if filename.strip() else ""
    assessment_note = (
        "Provide a full Questionnaire Assessment section critiquing alignment with the paper's main argument."
        if critique_enabled
        else "Assess whether it supports or helps answer the paper's main argument."
    )
    return (
        "\n\nResearch questionnaire / survey instrument related to the paper"
        f" ({assessment_note}):"
        f"{source_note}\n"
        "--- BEGIN QUESTIONNAIRE ---\n"
        f"{questionnaire_text.strip()}\n"
        "--- END QUESTIONNAIRE ---"
    )


def extract_optional_text_file(filename: str, data: bytes) -> str:
    if not filename or not data:
        return ""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported journal file type. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )
    return extract_text(filename, data)


def extract_optional_upload(filename: str, data: bytes, label: str) -> str:
    if not filename or not data:
        return ""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported {label} file type. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )
    return extract_text(filename, data)


_PAGE_NUMBER_RE = re.compile(r"^\s*(?:page\s+)?[-–—]?\s*\d{1,4}\s*(?:of\s+\d{1,4})?\s*[-–—]?\s*$", re.IGNORECASE)


def compress_text(text: str) -> str:
    """Reduce token footprint of extracted text without losing meaning.

    Removes PDF layout noise that wastes tokens (and triggers rate limits):
    repeated running headers/footers, standalone page numbers, hyphenated
    line breaks, and excess whitespace.
    """
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)  # rejoin hyphenated line breaks

    lines = text.split("\n")
    stripped = [line.strip() for line in lines]

    counts = Counter(line for line in stripped if line)
    cleaned: list[str] = []
    for line in stripped:
        if not line:
            cleaned.append("")
            continue
        if _PAGE_NUMBER_RE.match(line):
            continue
        # Exact-duplicate short lines repeated across many pages are almost
        # always running headers/footers (e.g. the paper title or author name).
        if len(line) <= 200 and counts[line] >= 5:
            continue
        cleaned.append(line)

    text = "\n".join(cleaned)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(data)
    elif ext in {".docx", ".doc"}:
        text = extract_text_from_docx(data)
    elif ext in {".txt", ".md"}:
        text = data.decode("utf-8", errors="replace")
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

    text = compress_text(text)
    if not text:
        raise HTTPException(
            status_code=400,
            detail="Could not extract readable text from this file. Try a different format or a text-based PDF.",
        )
    return text


def build_user_prompt(
    paper_text: str,
    filename: str,
    paper_type: str,
    review_depth: str,
    focus_areas: str,
    questioning_mode: str,
    reference_text: str = "",
    reference_link: str = "",
    journal_text: str = "",
    journal_filename: str = "",
    paper_truncated: bool = False,
    journal_truncated: bool = False,
    questionnaire_text: str = "",
    questionnaire_filename: str = "",
    questionnaire_truncated: bool = False,
    critique_questionnaire: bool = False,
    provider: str = "gemini",
) -> str:
    depth_instructions = {
        "overview": "Provide a concise review (roughly 800–1200 words). Focus on the most important issues.",
        "detailed": "Provide a thorough, section-by-section review (roughly 1500–2500 words). Be comprehensive.",
        "revision": "Focus heavily on actionable revision suggestions. Prioritize numbered edits the student can implement immediately.",
    }
    depth_note = depth_instructions.get(review_depth, depth_instructions["detailed"])

    focus_note = ""
    if focus_areas.strip():
        focus_note = f"\n\nThe author wants you to pay special attention to: {focus_areas.strip()}"

    truncation_note = (
        f"\n\n[Note: Only part of the paper was sent (~{len(paper_text):,} characters) due to "
        f"{'Groq free-tier' if provider == 'groq' else 'API'} limits. "
        "Review this excerpt and note if later sections may be missing.]"
        if paper_truncated
        else ""
    )
    reference_note = format_reference_block(reference_text, reference_link)
    journal_note = format_journal_block(journal_text, journal_filename)
    journal_truncation = (
        "\n\n[Note: The consultation journal was truncated due to length limits.]"
        if journal_truncated
        else ""
    )
    questionnaire_note = format_questionnaire_block(
        questionnaire_text,
        questionnaire_filename,
        critique_enabled=critique_questionnaire,
    )
    questionnaire_truncation = (
        "\n\n[Note: The questionnaire was truncated due to length limits.]"
        if questionnaire_truncated
        else ""
    )

    return f"""Please review the following academic document.

Document filename: {filename}
Paper type: {paper_type}
Review depth: {review_depth}
Questioning mode: {normalize_questioning_mode(questioning_mode)}
Questionnaire critique: {"enabled" if critique_questionnaire else "disabled" if not questionnaire_text else "alignment check only"}
Instructions: {depth_note}{focus_note}{truncation_note}{reference_note}{journal_note}{journal_truncation}{questionnaire_note}{questionnaire_truncation}

--- BEGIN DOCUMENT ---
{paper_text}
--- END DOCUMENT ---"""


app = FastAPI(title="Academic Paper Reviewer", version="1.0.0")


@app.get("/")
async def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


@app.post("/api/reference/load")
async def load_reference(
    drive_link: str = Form(default=""),
    ai_provider: str = Form(default="auto"),
    groq_api_key: str = Form(default=""),
    gemini_api_key: str = Form(default=""),
) -> dict:
    link = drive_link.strip()
    if not link:
        return build_reference_response("", "", False, None)

    provider = resolve_provider(ai_provider, groq_api_key, gemini_api_key)
    limits = get_limits(provider)
    text, truncated, warning = resolve_reference_material(
        link,
        max_reference_chars=limits.max_reference_chars,
    )
    return build_reference_response(link, text, truncated, warning)


@app.post("/api/review")
async def review_paper(
    file: UploadFile = File(...),
    api_key: str = Form(default=""),
    groq_api_key: str = Form(default=""),
    gemini_api_key: str = Form(default=""),
    ai_provider: str = Form(default="auto"),
    paper_type: str = Form(default="Research Paper"),
    review_depth: str = Form(default="detailed"),
    focus_areas: str = Form(default=""),
    questioning_mode: str = Form(default="balanced"),
    reference_link: str = Form(default=""),
    reference_excerpt: str = Form(default=""),
    journal_excerpt: str = Form(default=""),
    journal_filename: str = Form(default=""),
    journal_file: UploadFile | None = File(default=None),
    questionnaire_excerpt: str = Form(default=""),
    questionnaire_filename: str = Form(default=""),
    questionnaire_file: UploadFile | None = File(default=None),
    critique_questionnaire: str = Form(default="false"),
) -> dict:
    filename = file.filename or "document"
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    data = await file.read()
    if len(data) > MAX_FILE_BYTES:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 50 MB.")

    paper_text = extract_text(filename, data)
    if not paper_text.strip():
        raise HTTPException(status_code=400, detail="The uploaded paper contains no readable text.")

    critique_q = parse_form_bool(critique_questionnaire)
    groq_key = groq_api_key.strip() or api_key.strip()
    gemini_key = gemini_api_key.strip()
    provider = resolve_provider(ai_provider, groq_key, gemini_key)
    limits = get_limits(provider)

    questionnaire_text = questionnaire_excerpt.strip()
    questionnaire_name = questionnaire_filename.strip()
    if questionnaire_file and questionnaire_file.filename:
        questionnaire_data = await questionnaire_file.read()
        if len(questionnaire_data) > MAX_FILE_BYTES:
            raise HTTPException(status_code=400, detail="Questionnaire file too large. Maximum size is 50 MB.")
        questionnaire_name = questionnaire_file.filename
        questionnaire_text = extract_optional_upload(questionnaire_name, questionnaire_data, "questionnaire")

    has_questionnaire = bool(questionnaire_text.strip())
    if has_questionnaire and not paper_text.strip():
        raise HTTPException(
            status_code=400,
            detail="Upload the paper before submitting a questionnaire for assessment.",
        )
    if critique_q and not has_questionnaire:
        raise HTTPException(
            status_code=400,
            detail="Upload a questionnaire file to enable questionnaire critique.",
        )

    reference_text, reference_truncated, reference_warning = resolve_reference_material(
        reference_link,
        reference_excerpt,
        max_reference_chars=limits.max_reference_chars,
    )

    journal_text = journal_excerpt.strip()
    journal_name = journal_filename.strip()
    if journal_file and journal_file.filename:
        journal_data = await journal_file.read()
        if len(journal_data) > MAX_FILE_BYTES:
            raise HTTPException(status_code=400, detail="Consultation journal file too large. Maximum size is 50 MB.")
        journal_name = journal_file.filename
        journal_text = extract_optional_upload(journal_name, journal_data, "journal")

    review_system = get_review_system_prompt(
        provider,
        questioning_mode,
        critique_questionnaire=critique_q,
        include_questionnaire=has_questionnaire,
        paper_type=paper_type,
    )
    (
        paper_for_review,
        reference_text,
        journal_text,
        questionnaire_text,
        paper_truncated,
        ref_truncated_from_budget,
        journal_truncated,
        questionnaire_truncated,
    ) = fit_review_content(
        provider,
        paper_text,
        reference_text,
        journal_text,
        questionnaire_text,
        review_depth,
        review_system,
    )
    reference_truncated = reference_truncated or ref_truncated_from_budget

    user_prompt = build_user_prompt(
        paper_for_review,
        filename,
        paper_type,
        review_depth,
        focus_areas,
        questioning_mode,
        reference_text=reference_text,
        reference_link=reference_link.strip(),
        journal_text=journal_text,
        journal_filename=journal_name,
        paper_truncated=paper_truncated,
        journal_truncated=journal_truncated,
        questionnaire_text=questionnaire_text,
        questionnaire_filename=questionnaire_name,
        questionnaire_truncated=questionnaire_truncated,
        critique_questionnaire=critique_q,
        provider=provider,
    )

    review, model_name, provider_used = run_with_fallback(
        ai_provider,
        groq_key,
        gemini_key,
        task="review",
        system_prompt=review_system,
        user_prompt=user_prompt,
    )

    word_count = len(paper_text.split())

    return {
        "review": review,
        "filename": filename,
        "word_count": word_count,
        "model": model_name,
        "provider": provider_used,
        "truncated": paper_truncated,
        "paper_excerpt": paper_for_review[: limits.max_paper_excerpt],
        "paper_chars_sent": len(paper_for_review),
        "text_limit": get_text_limit(provider, review_depth),
        "reference_loaded": bool(reference_text),
        "reference_truncated": reference_truncated,
        "reference_excerpt": reference_text[:8000],
        "reference_word_count": len(reference_text.split()) if reference_text else 0,
        "reference_message": reference_warning,
        "journal_loaded": bool(journal_text),
        "journal_truncated": journal_truncated,
        "journal_excerpt": journal_text[: limits.max_journal_excerpt],
        "journal_filename": journal_name,
        "journal_word_count": len(journal_text.split()) if journal_text else 0,
        "questionnaire_loaded": bool(questionnaire_text),
        "questionnaire_critique": critique_q,
        "questionnaire_truncated": questionnaire_truncated,
        "questionnaire_excerpt": questionnaire_text[: limits.max_questionnaire_excerpt],
        "questionnaire_filename": questionnaire_name,
        "questionnaire_word_count": len(questionnaire_text.split()) if questionnaire_text else 0,
    }


CHAT_CONTEXT_PROMPT = """You are in a live conversation with a student about their academic writing.

CONVERSATION STYLE:
- Respond naturally and conversationally, like a supportive thesis advisor in office hours.
- Keep answers focused and practical — usually 2–6 paragraphs unless the student asks for detail.
- Reference the student's paper or your earlier review when that context is available.
- Ask clarifying questions when the student's request is vague.
- You may use markdown for lists or emphasis when helpful, but avoid the full formal review template unless asked.

REFERENCE MATERIAL (OPTIONAL):
- Google Drive reference material may or may not be available in this conversation.
- When it is missing or empty, rely on the uploaded paper, your earlier review, and the conversation history. Do not say Google Drive is required.

CONSULTATION JOURNAL (OPTIONAL):
- The student may provide a journal of meetings with other professors. Use it to align with prior advice and avoid repetition.
- When no journal is provided, rely on the paper, your review, and chat history only.

QUESTIONNAIRE (OPTIONAL):
- A research questionnaire may be provided only alongside the paper. Use it to discuss whether survey items support the thesis.
- When questionnaire critique was enabled in the review, reference that assessment in follow-up answers.

When questioning mode is active, treat unsupported or questionable student claims the way a rigorous advisor would — ask them to explain, defend, or revise."""


def parse_chat_history(raw: str) -> list[dict[str, str]]:
    if not raw.strip():
        return []
    try:
        data = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail="Invalid conversation history.") from exc
    if not isinstance(data, list):
        raise HTTPException(status_code=400, detail="Conversation history must be a list.")
    history: list[dict[str, str]] = []
    for item in data[-20:]:
        if not isinstance(item, dict):
            continue
        role = item.get("role")
        content = item.get("content")
        if role in {"user", "assistant"} and isinstance(content, str) and content.strip():
            history.append({"role": role, "content": content.strip()})
    return history


def build_chat_context(
    provider: str,
    filename: str,
    previous_review: str,
    paper_excerpt: str,
    reference_excerpt: str = "",
    reference_link: str = "",
    journal_excerpt: str = "",
    journal_filename: str = "",
    questionnaire_excerpt: str = "",
    questionnaire_filename: str = "",
    questionnaire_critique: bool = False,
) -> str:
    limits = get_limits(provider)
    parts: list[str] = []
    if filename.strip():
        parts.append(f"Document on file: {filename.strip()}")
    if previous_review.strip():
        parts.append(
            "Your earlier review of the student's paper:\n"
            + trim_context_text(previous_review, limits.max_review_excerpt)
        )
    if paper_excerpt.strip():
        parts.append(f"Paper excerpt:\n{trim_context_text(paper_excerpt, limits.max_paper_excerpt)}")
    if journal_excerpt.strip():
        label = journal_filename.strip() or "consultation journal"
        parts.append(
            f"Consultation journal ({label}) — notes from meetings with other professors:\n"
            + trim_context_text(journal_excerpt, limits.max_journal_excerpt)
        )
    if questionnaire_excerpt.strip():
        label = questionnaire_filename.strip() or "questionnaire"
        critique_note = (
            " (questionnaire critique was enabled in the review)"
            if questionnaire_critique
            else " — assess alignment with the paper's main argument"
        )
        parts.append(
            f"Research questionnaire ({label}){critique_note}:\n"
            + trim_context_text(questionnaire_excerpt, limits.max_questionnaire_excerpt)
        )
    if reference_excerpt.strip():
        source = f" ({reference_link.strip()})" if reference_link.strip() else ""
        parts.append(
            f"Google Drive reference material{source}:\n"
            + trim_context_text(reference_excerpt, limits.max_reference_chars)
        )
    if not parts:
        return ""
    return "Context for this conversation:\n\n" + "\n\n".join(parts)


def get_chat_system_prompt(provider: str, questioning_mode: str, paper_type: str = "") -> str:
    return (
        get_base_system_prompt(provider, include_template=should_include_template(paper_type))
        + "\n\n"
        + CHAT_CONTEXT_PROMPT
        + "\n\n"
        + get_questioning_prompt(questioning_mode, "chat")
    )


@app.post("/api/chat")
async def chat(
    api_key: str = Form(default=""),
    groq_api_key: str = Form(default=""),
    gemini_api_key: str = Form(default=""),
    ai_provider: str = Form(default="auto"),
    message: str = Form(...),
    history: str = Form(default="[]"),
    previous_review: str = Form(default=""),
    paper_excerpt: str = Form(default=""),
    filename: str = Form(default=""),
    paper_type: str = Form(default=""),
    questioning_mode: str = Form(default="balanced"),
    reference_link: str = Form(default=""),
    reference_excerpt: str = Form(default=""),
    journal_excerpt: str = Form(default=""),
    journal_filename: str = Form(default=""),
    questionnaire_excerpt: str = Form(default=""),
    questionnaire_filename: str = Form(default=""),
    questionnaire_critique: str = Form(default="false"),
) -> dict:
    if not message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty.")

    groq_key = groq_api_key.strip() or api_key.strip()
    gemini_key = gemini_api_key.strip()
    provider = resolve_provider(ai_provider, groq_key, gemini_key)
    prior = parse_chat_history(history)
    context = build_chat_context(
        provider,
        filename,
        previous_review,
        paper_excerpt,
        reference_excerpt=reference_excerpt,
        reference_link=reference_link,
        journal_excerpt=journal_excerpt,
        journal_filename=journal_filename,
        questionnaire_excerpt=questionnaire_excerpt,
        questionnaire_filename=questionnaire_filename,
        questionnaire_critique=parse_form_bool(questionnaire_critique),
    )
    chat_system = get_chat_system_prompt(provider, questioning_mode, paper_type=paper_type)

    response, model_name, provider_used = run_with_fallback(
        ai_provider,
        groq_key,
        gemini_key,
        task="chat",
        system_prompt=chat_system,
        context_prompt=context,
        history=prior,
        message=message.strip(),
    )
    return {"response": response, "provider": provider_used, "model": model_name}


@app.post("/api/follow-up")
async def follow_up(
    api_key: str = Form(default=""),
    groq_api_key: str = Form(default=""),
    gemini_api_key: str = Form(default=""),
    ai_provider: str = Form(default="auto"),
    question: str = Form(...),
    paper_excerpt: str = Form(default=""),
    previous_review: str = Form(default=""),
    paper_type: str = Form(default=""),
    questioning_mode: str = Form(default="balanced"),
    reference_link: str = Form(default=""),
    reference_excerpt: str = Form(default=""),
    journal_excerpt: str = Form(default=""),
    journal_filename: str = Form(default=""),
    questionnaire_excerpt: str = Form(default=""),
    questionnaire_filename: str = Form(default=""),
    questionnaire_critique: str = Form(default="false"),
) -> dict:
    if not question.strip():
        raise HTTPException(status_code=400, detail="Question cannot be empty.")

    groq_key = groq_api_key.strip() or api_key.strip()
    gemini_key = gemini_api_key.strip()
    provider = resolve_provider(ai_provider, groq_key, gemini_key)
    context = build_chat_context(
        provider,
        "",
        previous_review,
        paper_excerpt,
        reference_excerpt=reference_excerpt,
        reference_link=reference_link,
        journal_excerpt=journal_excerpt,
        journal_filename=journal_filename,
        questionnaire_excerpt=questionnaire_excerpt,
        questionnaire_filename=questionnaire_filename,
        questionnaire_critique=parse_form_bool(questionnaire_critique),
    )
    chat_system = get_chat_system_prompt(provider, questioning_mode, paper_type=paper_type)

    response, model_name, provider_used = run_with_fallback(
        ai_provider,
        groq_key,
        gemini_key,
        task="chat",
        system_prompt=chat_system,
        context_prompt=context,
        history=[],
        message=question.strip(),
    )
    return {"response": response, "provider": provider_used, "model": model_name}
